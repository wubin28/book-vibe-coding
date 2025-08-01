import sys
import os
import re
import shutil
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def read_markdown_file(file_path):
    """Read the content of a Markdown file."""
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()

def parse_markdown(markdown_text):
    """Parse Markdown text into structured elements for Word conversion."""
    lines = markdown_text.split('\n')
    elements = []
    
    i = 0
    in_code_block = False
    code_block_lines = []
    code_block_info = ""
    current_paragraph_lines = []
    in_aside = False
    aside_lines = []
    
    while i < len(lines):
        line = lines[i]
        
        # Handle code blocks
        if line.startswith('```'):
            if not in_code_block:
                # Start of code block
                in_code_block = True
                if current_paragraph_lines:
                    elements.append(('paragraph', '\n'.join(current_paragraph_lines)))
                    current_paragraph_lines = []
                code_block_info = line[3:].strip()
                code_block_lines = []
            else:
                # End of code block
                in_code_block = False
                elements.append(('code_block', code_block_info, '\n'.join(code_block_lines)))
            i += 1
            continue
        
        if in_code_block:
            code_block_lines.append(line)
            i += 1
            continue
        
        # Handle asides
        if line.strip() == '<aside>':
            in_aside = True
            if current_paragraph_lines:
                elements.append(('paragraph', '\n'.join(current_paragraph_lines)))
                current_paragraph_lines = []
            i += 1
            continue
        
        if line.strip() == '</aside>':
            in_aside = False
            elements.append(('aside', '\n'.join(aside_lines)))
            aside_lines = []
            i += 1
            continue
        
        if in_aside:
            aside_lines.append(line)
            i += 1
            continue
        
        # Handle headings
        heading_match = re.match(r'^(#+)\s+(.*)', line)
        if heading_match:
            if current_paragraph_lines:
                elements.append(('paragraph', '\n'.join(current_paragraph_lines)))
                current_paragraph_lines = []
            
            level = len(heading_match.group(1))
            heading_text = heading_match.group(2)
            elements.append(('heading', level, heading_text))
            i += 1
            continue
        
        # Handle images
        image_match = re.match(r'!\[(.*?)\]\((.*?)\)', line)
        if image_match:
            if current_paragraph_lines:
                elements.append(('paragraph', '\n'.join(current_paragraph_lines)))
                current_paragraph_lines = []
            
            alt_text = image_match.group(1)
            image_path = image_match.group(2)
            
            # Check for figure caption on next line
            caption = None
            if i + 1 < len(lines) and (lines[i + 1].strip().startswith('图') or "图" in lines[i + 1]):
                caption = lines[i + 1].strip()
                i += 2  # Skip both the image and caption lines
            else:
                i += 1  # Skip just the image line
            
            elements.append(('image', alt_text, image_path, caption))
            continue
        
        # Handle bullet lists
        bullet_list_match = re.match(r'^(\s*)[-*]\s+(.*)', line)
        if bullet_list_match:
            if current_paragraph_lines:
                elements.append(('paragraph', '\n'.join(current_paragraph_lines)))
                current_paragraph_lines = []
            
            indent = len(bullet_list_match.group(1))
            item_text = bullet_list_match.group(2)
            elements.append(('bullet_list', indent, item_text))
            i += 1
            continue
        
        # Handle numbered lists
        numbered_list_match = re.match(r'^(\s*)(\d+)[.)]\s+(.*)', line)
        if numbered_list_match:
            if current_paragraph_lines:
                elements.append(('paragraph', '\n'.join(current_paragraph_lines)))
                current_paragraph_lines = []
            
            indent = len(numbered_list_match.group(1))
            number = numbered_list_match.group(2)
            item_text = numbered_list_match.group(3)
            elements.append(('numbered_list', indent, number, item_text))
            i += 1
            continue
        
        # Handle special sections like 【避坑指南】
        if "【避坑指南】" in line:
            if current_paragraph_lines:
                elements.append(('paragraph', '\n'.join(current_paragraph_lines)))
                current_paragraph_lines = []
            elements.append(('special_section', line))
            i += 1
            continue
            
        # Handle code listings
        code_listing_match = re.match(r'^代码清单(\d+-\d+)[\s]*(.*)', line)
        if code_listing_match:
            if current_paragraph_lines:
                elements.append(('paragraph', '\n'.join(current_paragraph_lines)))
                current_paragraph_lines = []
            
            listing_number = code_listing_match.group(1)
            listing_title = code_listing_match.group(2)
            elements.append(('code_listing', listing_number, listing_title))
            i += 1
            continue
        
        # Handle empty lines
        if not line.strip():
            if current_paragraph_lines:
                elements.append(('paragraph', '\n'.join(current_paragraph_lines)))
                current_paragraph_lines = []
            i += 1
            continue
        
        # Default: add to current paragraph
        current_paragraph_lines.append(line)
        i += 1
    
    # Add any remaining paragraph content
    if current_paragraph_lines:
        elements.append(('paragraph', '\n'.join(current_paragraph_lines)))
    
    return elements

def create_docx_from_elements(elements, template_path, output_path):
    """Create a Word document from parsed Markdown elements."""
    # Copy the template file
    shutil.copy2(template_path, output_path)
    
    # Open the copied template
    doc = Document(output_path)
    
    # Clear existing content while preserving styles
    for i in range(len(doc.paragraphs)-1, -1, -1):
        p = doc.paragraphs[i]._element
        p.getparent().remove(p)
    
    # Process elements and add to document
    for element in elements:
        element_type = element[0]
        
        if element_type == 'heading':
            level, text = element[1], element[2]
            if level == 1:
                heading = doc.add_heading('', 0)  # Title style for chapter titles
            else:
                heading = doc.add_heading('', level)
            heading.add_run(text)
        
        elif element_type == 'paragraph':
            text = element[1]
            para = doc.add_paragraph()
            para.add_run(text)
        
        elif element_type == 'code_block':
            language, code = element[1], element[2]
            
            # Check if this is a code listing with a title embedded in the content
            if code.startswith('代码清单'):
                lines = code.split('\n')
                title_line = lines[0]
                code_content = '\n'.join(lines[1:])
                
                # Add the title
                title_para = doc.add_paragraph()
                title_para.add_run(title_line).bold = True
                
                # Add the code content
                code_para = doc.add_paragraph()
                code_run = code_para.add_run(code_content)
                code_run.font.name = 'Courier New'
                code_run.font.size = Pt(9)
            else:
                # Regular code block
                code_para = doc.add_paragraph()
                code_run = code_para.add_run(code)
                code_run.font.name = 'Courier New'
                code_run.font.size = Pt(9)
        
        elif element_type == 'code_listing':
            number, title = element[1], element[2]
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(f"代码清单{number} {title}").bold = True
            
            # If there's no title, just use the number
            if not title.strip():
                title_para.add_run(f"代码清单{number}").bold = True
        
        elif element_type == 'image':
            alt_text, image_path = element[1], element[2]
            caption = element[3] if len(element) > 3 else None
            
            # Add placeholder for image (we won't insert actual images)
            img_para = doc.add_paragraph()
            img_para.add_run(f"![{alt_text}]({image_path})")
            
            # Add caption if available
            if caption:
                caption_para = doc.add_paragraph()
                caption_para.add_run(caption)
                caption_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        elif element_type == 'bullet_list':
            indent, text = element[1], element[2]
            list_para = doc.add_paragraph(style='List Bullet')
            list_para.add_run(text)
        
        elif element_type == 'numbered_list':
            indent, number, text = element[1], element[2], element[3]
            list_para = doc.add_paragraph(style='List Number')
            list_para.add_run(text)
        
        elif element_type == 'aside':
            text = element[1]
            # Handle special formatting for asides
            aside_para = doc.add_paragraph()
            aside_run = aside_para.add_run(text)
            aside_run.italic = True
            
        elif element_type == 'special_section':
            # Special handling for sections like 【避坑指南】
            text = element[1]
            special_para = doc.add_paragraph()
            special_run = special_para.add_run(text)
            special_run.bold = True
    
    # Save the document
    doc.save(output_path)

def main():
    """Main function to execute the conversion."""
    if len(sys.argv) != 2:
        print("Usage: python converter-by-claude-sonnet-3.7-with-old-prompt.py <markdown_file>")
        sys.exit(1)
    
    markdown_file = sys.argv[1]
    template_file = "ch04-to-template.docx"
    output_file = "ch04-to.docx"
    
    # Check if files exist
    if not os.path.exists(markdown_file):
        print(f"Error: Markdown file '{markdown_file}' not found.")
        sys.exit(1)
    
    if not os.path.exists(template_file):
        print(f"Error: Template file '{template_file}' not found.")
        sys.exit(1)
    
    # Read and parse the Markdown file
    print(f"Reading Markdown file: {markdown_file}")
    markdown_text = read_markdown_file(markdown_file)
    
    print("Parsing Markdown content...")
    elements = parse_markdown(markdown_text)
    
    print(f"Creating Word document: {output_file}")
    create_docx_from_elements(elements, template_file, output_file)
    
    print(f"Conversion completed. Output file: {output_file}")
    print("Note: Images are not included in the converted document.")

if __name__ == "__main__":
    main()